import os
import sys
import json
import asyncio
import random
import tempfile
import subprocess
import re
import time
import base64
from datetime import datetime, time as dt_time
from zoneinfo import ZoneInfo
from io import BytesIO
from telethon import TelegramClient, events
from telethon.sessions import StringSession
from anthropic import Anthropic
from supabase import create_client, Client
import PyPDF2
from docx import Document
import gspread
from google.oauth2.service_account import Credentials

# Add shared directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from shared.knowledgebase import (
    RECRUITER_NAME, COMPANY_NAME, APPLICATION_FORM_URL,
    ConversationContext, build_system_prompt, build_context_from_state,
    identify_role_from_text, identify_role_semantic, get_experience_question,
    get_resume_acknowledgment, reload_from_database, embed_existing_knowledge
)
from shared.training_handlers import handle_training_message, init_admin_users
from shared.database import (
    save_message, get_conversation_messages, get_or_create_conversation_state,
    update_conversation_state_db, link_conversation_to_candidate
)

# Conversation memory (max messages per user for AI context)
conversations = {}
MAX_MESSAGES = 25  # Increased for better context
conversation_states = {}

# Track which users have been restored from DB
restored_users = set()

# Operating hours configuration (Singapore timezone)
TIMEZONE = ZoneInfo("Asia/Singapore")
OPERATING_START = dt_time(8, 30)  # 8:30 AM
OPERATING_END = dt_time(22, 0)    # 10:00 PM

# Knowledgebase auto-refresh interval (5 minutes)
KB_REFRESH_INTERVAL = 300  # seconds


async def periodic_knowledgebase_refresh():
    """
    Background task that periodically refreshes the knowledgebase from database.
    This ensures the bot picks up changes made in the CRM without requiring a restart.
    """
    while True:
        try:
            await asyncio.sleep(KB_REFRESH_INTERVAL)
            print(f"[{datetime.now(TIMEZONE)}] Auto-refreshing knowledgebase...")
            success = await reload_from_database()
            if success:
                print(f"[{datetime.now(TIMEZONE)}] Knowledgebase auto-refresh complete")
            else:
                print(f"[{datetime.now(TIMEZONE)}] Knowledgebase auto-refresh: no changes")
        except Exception as e:
            print(f"[{datetime.now(TIMEZONE)}] Error during knowledgebase auto-refresh: {e}")


def is_within_operating_hours() -> bool:
    """Check if current time is within operating hours (8:30 AM - 10:00 PM Singapore time)."""
    now = datetime.now(TIMEZONE)
    current_time = now.time()
    return OPERATING_START <= current_time <= OPERATING_END


# Spam protection
rate_limit_tracker = {}  # {user_id: [timestamp1, timestamp2, ...]}
RATE_LIMIT_MESSAGES = 10  # Max messages per time window
RATE_LIMIT_WINDOW = 60  # Time window in seconds (1 minute)

# Spam keywords to ignore (case-insensitive)
SPAM_KEYWORDS = [
    "crypto", "bitcoin", "ethereum", "investment opportunity",
    "make money fast", "work from home", "earn $", "earn usd",
    "click here", "free money", "lottery", "you have won",
    "nigerian prince", "wire transfer", "western union",
    "telegram premium", "free premium", "hack", "password"
]


def get_blocked_users() -> set:
    """Get set of blocked user IDs from environment variable."""
    blocked = os.environ.get('BLOCKED_TELEGRAM_USERS', '')
    if not blocked:
        return set()
    try:
        return set(int(uid.strip()) for uid in blocked.split(',') if uid.strip())
    except ValueError:
        return set()


def get_whitelist_users() -> set:
    """Get set of whitelisted user IDs (if whitelist mode is enabled)."""
    whitelist = os.environ.get('WHITELIST_TELEGRAM_USERS', '')
    if not whitelist:
        return set()
    try:
        return set(int(uid.strip()) for uid in whitelist.split(',') if uid.strip())
    except ValueError:
        return set()


def is_whitelist_mode() -> bool:
    """Check if whitelist mode is enabled."""
    return os.environ.get('TELEGRAM_WHITELIST_MODE', '').lower() in ('true', '1', 'yes')


def is_user_allowed(user_id: int) -> tuple[bool, str]:
    """Check if a user is allowed to interact with the bot.
    Returns (allowed, reason) tuple.
    """
    # Check blocked list
    if user_id in get_blocked_users():
        return False, "blocked"

    # Check whitelist mode
    if is_whitelist_mode():
        whitelist = get_whitelist_users()
        if whitelist and user_id not in whitelist:
            return False, "not_whitelisted"

    return True, "allowed"


def is_rate_limited(user_id: int) -> bool:
    """Check if user has exceeded rate limit."""
    current_time = time.time()

    if user_id not in rate_limit_tracker:
        rate_limit_tracker[user_id] = []

    # Remove old timestamps outside the window
    rate_limit_tracker[user_id] = [
        ts for ts in rate_limit_tracker[user_id]
        if current_time - ts < RATE_LIMIT_WINDOW
    ]

    # Check if over limit
    if len(rate_limit_tracker[user_id]) >= RATE_LIMIT_MESSAGES:
        return True

    # Add current timestamp
    rate_limit_tracker[user_id].append(current_time)
    return False


def contains_spam(text: str) -> bool:
    """Check if message contains spam keywords."""
    if not text:
        return False
    text_lower = text.lower()
    return any(keyword in text_lower for keyword in SPAM_KEYWORDS)


async def check_spam_protection(event, user_id: int, username: str, text: str = "") -> bool:
    """Run all spam checks. Returns True if message should be ignored."""
    # Check if user is allowed
    allowed, reason = is_user_allowed(user_id)
    if not allowed:
        print(f"Blocked user {user_id} (@{username}): {reason}")
        return True

    # Check rate limit
    if is_rate_limited(user_id):
        print(f"Rate limited user {user_id} (@{username})")
        # Optionally send a warning (only once)
        if len(rate_limit_tracker.get(user_id, [])) == RATE_LIMIT_MESSAGES:
            await event.respond("You're sending messages too quickly. Please wait a moment.")
        return True

    # Check for spam content
    if contains_spam(text):
        print(f"Spam detected from {user_id} (@{username}): {text[:50]}...")
        return True

    return False

# Conversation state tracking for dynamic prompts
conversation_states = {}

# Resume screening prompt (kept local for Telegram bot)
SCREENING_PROMPT = """You are analyzing a resume for a staffing agency. Your task is to evaluate the candidate.

AVAILABLE JOB ROLES:
{job_roles}

RESUME TEXT:
{resume_text}

INSTRUCTIONS:
1. Identify which job role the candidate is applying for based on context. Match to one of the available roles.
2. Analyze the resume against that role's requirements.
3. Extract contact information (email, phone) if visible.

CITIZENSHIP REQUIREMENT (CRITICAL):
Most roles require Singapore Citizens or Permanent Residents. Look for indicators:
- NRIC number (S/T = Citizen, F/G = PR)
- National Service (NS) completion
- Singapore address
- Local education (NUS, NTU, SMU, polytechnics, ITE)
- Explicit mention of citizenship/PR status

If no clear indicator of SC/PR status is found, set recommendation to "Rejected" unless the role specifically allows foreigners.

RESPONSE FORMAT:
Return a JSON object:
{{
    "candidate_name": "Full name from resume",
    "candidate_email": "email@example.com or null",
    "candidate_phone": "+65 XXXX XXXX or null",
    "job_matched": "Best matching role from your list",
    "score": 7,
    "citizenship_status": "Singapore Citizen|PR|Unknown|Foreigner",
    "recommendation": "Top Candidate|Review|Rejected",
    "summary": "Brief evaluation including citizenship verification"
}}

Use the scoring guide for the matched role. Score 1-10."""

# Global clients - initialized in main()
client = None
anthropic_client = None
supabase_client = None
gsheets_client = None
job_roles_cache = None
job_roles_cache_time = 0
JOB_ROLES_CACHE_DURATION = 300  # 5 minutes


async def restore_conversation_from_db(user_id: int):
    """Restore conversation history and state from database for returning users."""
    already_restored = user_id in restored_users

    try:
        # Always ensure conversation state exists (creates if deleted from DB)
        db_state = await get_or_create_conversation_state("telegram", str(user_id))
        if db_state:
            conversation_states[user_id] = {
                "stage": db_state.get("stage", "new"),
                "applied_role": db_state.get("applied_role"),
                "candidate_name": db_state.get("candidate_name"),
                "resume_received": db_state.get("resume_received", False),
                "form_completed": db_state.get("form_completed", False),
                "experience_discussed": db_state.get("experience_discussed", False),
                "citizenship_status": db_state.get("citizenship_status"),
            }

        # Only restore messages once (avoid duplicates in memory)
        if not already_restored:
            # Get messages from database
            db_messages = await get_conversation_messages("telegram", str(user_id), limit=MAX_MESSAGES * 2)

            if db_messages:
                # Convert to the format we use in memory
                conversations[user_id] = [
                    {"role": msg["role"], "content": msg["content"]}
                    for msg in db_messages
                ]
                print(f"Restored {len(db_messages)} messages for user {user_id}")

            restored_users.add(user_id)

    except Exception as e:
        print(f"Error restoring conversation from DB: {e}")
        restored_users.add(user_id)  # Mark as restored to avoid repeated errors


def get_conversation(user_id: int) -> list:
    if user_id not in conversations:
        conversations[user_id] = []
    return conversations[user_id]


async def add_message_async(user_id: int, role: str, content: str):
    """Add message to memory and save to database."""
    conv = get_conversation(user_id)
    conv.append({"role": role, "content": content})
    if len(conv) > MAX_MESSAGES * 2:
        conv[:] = conv[-MAX_MESSAGES * 2:]

    # Save to database asynchronously
    try:
        await save_message("telegram", str(user_id), role, content)
    except Exception as e:
        print(f"Error saving message to DB: {e}")


def add_message(user_id: int, role: str, content: str):
    """Sync wrapper - adds to memory only. Use add_message_async for DB persistence."""
    conv = get_conversation(user_id)
    conv.append({"role": role, "content": content})
    if len(conv) > MAX_MESSAGES * 2:
        conv[:] = conv[-MAX_MESSAGES * 2:]


def get_conversation_state(user_id: int) -> dict:
    """Get the conversation state for a user."""
    if user_id not in conversation_states:
        conversation_states[user_id] = {
            "stage": "new",
            "applied_role": None,
            "candidate_name": None,
            "resume_received": False,
            "form_completed": False,
            "experience_discussed": False,
            "citizenship_status": None,
            "call_scheduled": False
        }
    return conversation_states[user_id]


def update_conversation_state(user_id: int, **kwargs):
    """Update the conversation state for a user (memory only)."""
    state = get_conversation_state(user_id)
    state.update(kwargs)
    return state


async def update_conversation_state_async(user_id: int, **kwargs):
    """Update the conversation state for a user and persist to database."""
    state = get_conversation_state(user_id)
    state.update(kwargs)

    # Save to database
    try:
        await update_conversation_state_db("telegram", str(user_id), **kwargs)
    except Exception as e:
        print(f"Error saving conversation state to DB: {e}")

    return state


async def detect_state_from_message_async(user_id: int, message: str):
    """Detect and update state based on message content, persisting to DB."""
    state = get_conversation_state(user_id)
    message_lower = message.lower()

    # Detect form completion
    form_keywords = ["done", "completed", "finished", "filled", "submitted"]
    if any(kw in message_lower for kw in form_keywords):
        if not state["form_completed"]:
            await update_conversation_state_async(user_id, form_completed=True, stage="form_completed")

    # Detect job role using semantic search (RAG) with keyword fallback
    detected_role = await identify_role_semantic(message)
    if detected_role and detected_role != "general":
        await update_conversation_state_async(user_id, applied_role=detected_role)

    # Detect citizenship mentions
    citizenship_map = {
        "singapore citizen": "SC", "singaporean": "SC", "sg citizen": "SC",
        "permanent resident": "PR", " pr ": "PR", "foreigner": "Foreign"
    }
    for indicator, status in citizenship_map.items():
        if indicator in message_lower:
            await update_conversation_state_async(user_id, citizenship_status=status)
            break


def detect_state_from_message(user_id: int, message: str):
    """Detect and update state based on message content (memory only)."""
    state = get_conversation_state(user_id)
    message_lower = message.lower()

    # Detect form completion
    form_keywords = ["done", "completed", "finished", "filled", "submitted"]
    if any(kw in message_lower for kw in form_keywords):
        if not state["form_completed"]:
            update_conversation_state(user_id, form_completed=True, stage="form_completed")

    # Detect job role
    detected_role = identify_role_from_text(message)
    if detected_role and detected_role != "general":
        update_conversation_state(user_id, applied_role=detected_role)

    # Detect citizenship mentions
    citizenship_map = {
        "singapore citizen": "SC", "singaporean": "SC", "sg citizen": "SC",
        "permanent resident": "PR", " pr ": "PR", "foreigner": "Foreign"
    }
    for indicator, status in citizenship_map.items():
        if indicator in message_lower:
            update_conversation_state(user_id, citizenship_status=status)
            break


async def get_ai_response(user_id: int, message: str, candidate_name: str = None) -> str:
    """Get AI response using dynamic context-aware prompting with RAG."""
    # Restore conversation from database if this is a returning user
    await restore_conversation_from_db(user_id)

    # Ensure message is not empty
    if not message or not message.strip():
        message = "[Empty message]"

    # Detect and update state based on message (with DB persistence)
    await detect_state_from_message_async(user_id, message)

    # Store candidate name if provided and not already set
    current_state = get_conversation_state(user_id)
    if candidate_name and not current_state.get("candidate_name"):
        await update_conversation_state_async(user_id, candidate_name=candidate_name)

    # Add user message and save to database
    await add_message_async(user_id, "user", message)

    # Filter out any empty messages from conversation history
    valid_messages = [
        msg for msg in get_conversation(user_id)
        if msg.get("content") and msg["content"].strip()
    ]

    if not valid_messages:
        # First message - simple greeting with form
        greeting = f"Hi! I'm {RECRUITER_NAME} from {COMPANY_NAME} :)\n---\nCould u fill up this form? {APPLICATION_FORM_URL}\n---\nJust select '{RECRUITER_NAME}' as the consultant"
        await add_message_async(user_id, "assistant", greeting)
        return greeting

    # Build dynamic context-aware system prompt
    state = get_conversation_state(user_id)
    context = build_context_from_state(str(user_id), state)
    system_prompt = build_system_prompt(context)

    # RAG: Retrieve relevant context from knowledgebase
    try:
        from shared.knowledgebase import get_relevant_context_for_query
        rag_context = await get_relevant_context_for_query(message)
        if rag_context:
            system_prompt = system_prompt + "\n" + rag_context
    except Exception as e:
        print(f"RAG context retrieval failed (continuing without): {e}")

    try:
        response = anthropic_client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=1024,
            system=system_prompt,
            messages=valid_messages
        )
        ai_message = response.content[0].text

        # Save assistant response to database
        await add_message_async(user_id, "assistant", ai_message)

        # Update state based on response and persist to DB
        response_lower = ai_message.lower()
        if "resume" in response_lower and ("send" in response_lower or "have" in response_lower):
            if state.get("form_completed"):
                await update_conversation_state_async(user_id, stage="resume_requested")
        if any(phrase in response_lower for phrase in ["will contact", "shortlisted"]):
            await update_conversation_state_async(user_id, stage="conversation_closed")

        return ai_message
    except Exception as e:
        print(f"Error getting AI response: {e}")
        return "sorry, having some trouble. could u try again?"


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text content from a PDF file."""
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return ""


async def extract_text_from_pdf_with_vision(pdf_bytes: bytes) -> str:
    """Extract text from a PDF using Claude's vision API (for image-based PDFs like Canva resumes).

    This is used as a fallback when PyPDF2 cannot extract text (e.g., when text is rendered as images).
    """
    global anthropic_client

    if not anthropic_client:
        print("Error: Anthropic client not initialized for vision extraction")
        return ""

    try:
        # Encode PDF as base64
        pdf_base64 = base64.standard_b64encode(pdf_bytes).decode("utf-8")

        print("Using Claude vision API to extract text from image-based PDF...")

        response = anthropic_client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "document",
                            "source": {
                                "type": "base64",
                                "media_type": "application/pdf",
                                "data": pdf_base64
                            }
                        },
                        {
                            "type": "text",
                            "text": """Extract ALL text content from this resume/CV document.
Include everything: name, contact details, work experience, education, skills, certifications, etc.
Format it in a readable way, preserving the structure and sections.
Just output the extracted text, no commentary."""
                        }
                    ]
                }
            ]
        )

        extracted_text = response.content[0].text
        print(f"Vision API extracted {len(extracted_text)} characters from PDF")
        return extracted_text.strip()

    except Exception as e:
        print(f"Error extracting PDF text with vision API: {e}")
        return ""


def extract_text_from_word(doc_bytes: bytes) -> str:
    """Extract text content from a Word document (.docx)."""
    try:
        doc = Document(BytesIO(doc_bytes))
        text_parts = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"Error extracting Word document text: {e}")
        return ""


def convert_word_to_pdf(doc_bytes: bytes) -> bytes:
    """Convert a Word document to PDF using LibreOffice."""
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Write Word doc to temp file
            input_path = os.path.join(tmp_dir, "input.docx")
            with open(input_path, "wb") as f:
                f.write(doc_bytes)

            # Convert using LibreOffice
            subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", tmp_dir,
                input_path
            ], check=True, capture_output=True, timeout=60)

            # Read the converted PDF
            output_path = os.path.join(tmp_dir, "input.pdf")
            with open(output_path, "rb") as f:
                return f.read()
    except subprocess.TimeoutExpired:
        print("Error: Word to PDF conversion timed out")
        return None
    except Exception as e:
        print(f"Error converting Word to PDF: {e}")
        return None


def get_job_roles_from_sheets() -> str:
    """Fetch job roles from Google Sheets with caching."""
    global job_roles_cache, job_roles_cache_time, gsheets_client

    import time
    current_time = time.time()

    # Return cached data if still valid
    if job_roles_cache and (current_time - job_roles_cache_time) < JOB_ROLES_CACHE_DURATION:
        return job_roles_cache

    if not gsheets_client:
        print("Warning: Google Sheets not configured, using default job roles")
        return "No specific job roles configured. Screen the resume generally."

    try:
        spreadsheet_id = os.environ.get('GOOGLE_SHEETS_ID')
        if not spreadsheet_id:
            print("Warning: GOOGLE_SHEETS_ID not set")
            return "No specific job roles configured."

        sheet = gsheets_client.open_by_key(spreadsheet_id).sheet1
        records = sheet.get_all_records()

        job_roles_text = ""
        for row in records:
            job_title = row.get('Job Title', '')
            requirements = row.get('Requirements', '')
            scoring = row.get('Scoring Guide', '')
            if job_title:
                job_roles_text += f"\n\nJOB: {job_title}\nRequirements: {requirements}\nScoring: {scoring}"

        job_roles_cache = job_roles_text
        job_roles_cache_time = current_time
        print(f"Fetched {len(records)} job roles from Google Sheets")
        return job_roles_text

    except Exception as e:
        print(f"Error fetching job roles from Google Sheets: {e}")
        return job_roles_cache or "No specific job roles configured."


async def screen_resume(resume_text: str) -> dict:
    """Use AI to screen the resume against job requirements."""
    try:
        job_roles = get_job_roles_from_sheets()

        prompt = SCREENING_PROMPT.format(
            job_roles=job_roles,
            resume_text=resume_text[:15000]  # Limit resume text length
        )

        response = anthropic_client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=2048,
            messages=[{"role": "user", "content": prompt}]
        )

        response_text = response.content[0].text

        # Try to parse JSON from response
        try:
            # Find JSON in response
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                return json.loads(json_match.group())
        except json.JSONDecodeError:
            pass

        # Return basic structure if parsing fails
        return {
            "matched_job": "Unknown",
            "score": 5,
            "qualifications": [],
            "missing": [],
            "recommendation": "Review",
            "reason": response_text[:500],
            "raw_response": response_text
        }

    except Exception as e:
        print(f"Error screening resume: {e}")
        return {
            "error": str(e),
            "recommendation": "Review",
            "reason": "Screening failed - manual review required"
        }


async def upload_resume_to_storage(file_bytes: bytes, file_name: str, user_id: int) -> str:
    """Upload resume file to Supabase Storage and return the public URL."""
    try:
        import time
        # Create a unique filename
        timestamp = int(time.time())
        safe_name = file_name.replace(' ', '_')
        storage_path = f"resumes/{user_id}_{timestamp}_{safe_name}"

        # Upload to Supabase Storage
        result = supabase_client.storage.from_("resumes").upload(
            storage_path,
            file_bytes,
            {"content-type": "application/pdf"}
        )

        # Get public URL
        public_url = supabase_client.storage.from_("resumes").get_public_url(storage_path)
        print(f"Resume uploaded to: {public_url}")
        return public_url

    except Exception as e:
        print(f"Error uploading resume to storage: {e}")
        # Try creating the bucket if it doesn't exist
        try:
            supabase_client.storage.create_bucket("resumes", {"public": True})
            # Retry upload
            result = supabase_client.storage.from_("resumes").upload(
                storage_path,
                file_bytes,
                {"content-type": "application/pdf"}
            )
            public_url = supabase_client.storage.from_("resumes").get_public_url(storage_path)
            print(f"Resume uploaded to: {public_url}")
            return public_url
        except Exception as e2:
            print(f"Failed to create bucket and upload: {e2}")
            return None


async def save_candidate(user_id: int, username: str, full_name: str, screening_result: dict = None, resume_url: str = None):
    """Save or update candidate in database with optional screening results."""
    try:
        conv_history = get_conversation(user_id)

        data = {
            "full_name": full_name or f"Telegram User {user_id}",
            "telegram_user_id": user_id,
            "telegram_username": username,
            "source": "telegram",
            "conversation_history": json.dumps(conv_history)
        }

        # Add screening results if available
        if screening_result:
            # Update name/email/phone if extracted from resume
            if screening_result.get("candidate_name"):
                data["full_name"] = screening_result["candidate_name"]
            if screening_result.get("candidate_email"):
                data["email"] = screening_result["candidate_email"]
            if screening_result.get("candidate_phone"):
                data["phone"] = screening_result["candidate_phone"]

            # Map recommendation to ai_category
            rec = screening_result.get("recommendation", "Review")
            if "Top" in rec:
                data["ai_category"] = "Top Candidate"
            elif "Reject" in rec:
                data["ai_category"] = "Rejected"
            else:
                data["ai_category"] = "Review"

            # Set current_status to ai_screened when resume is processed
            data["current_status"] = "ai_screened"

            # Add screening data
            data["applied_role"] = screening_result.get("job_matched", "")
            data["ai_score"] = screening_result.get("score", 0)
            data["ai_summary"] = screening_result.get("summary", "")

            # Add citizenship status - map to database format
            citizenship = screening_result.get("citizenship_status", "")
            if citizenship == "Singapore Citizen":
                data["citizenship_status"] = "SC"
            elif citizenship == "PR":
                data["citizenship_status"] = "PR"
            elif citizenship == "Foreigner":
                data["citizenship_status"] = "Foreign"
            else:
                data["citizenship_status"] = "Not Identified"

            # Store full screening result as JSON if column exists
            try:
                data["screening_result"] = json.dumps(screening_result)
            except:
                pass
        else:
            data["current_status"] = "new_application"

        # Add resume URL if provided
        if resume_url:
            data["resume_url"] = resume_url

        existing = supabase_client.table("candidates").select("id").eq("telegram_user_id", user_id).execute()
        if existing.data:
            supabase_client.table("candidates").update(data).eq("telegram_user_id", user_id).execute()
            print(f"Updated candidate: {data['full_name']}")
        else:
            supabase_client.table("candidates").insert(data).execute()
            print(f"Created new candidate: {data['full_name']}")

        return True
    except Exception as e:
        print(f"Error saving candidate: {e}")
        return False


def validate_env_vars():
    """Validate all required environment variables are set."""
    required = {
        'TELEGRAM_API_ID': os.environ.get('TELEGRAM_API_ID'),
        'TELEGRAM_API_HASH': os.environ.get('TELEGRAM_API_HASH'),
        'TELEGRAM_SESSION_STRING': os.environ.get('TELEGRAM_SESSION_STRING'),
        'CLAUDE_API_KEY': os.environ.get('CLAUDE_API_KEY'),
        'SUPABASE_URL': os.environ.get('SUPABASE_URL'),
        'SUPABASE_ANON_KEY': os.environ.get('SUPABASE_ANON_KEY'),
    }

    missing = [key for key, value in required.items() if not value]

    if missing:
        print(f"ERROR: Missing required environment variables: {', '.join(missing)}")
        return False

    # Validate API_ID is a number
    try:
        int(required['TELEGRAM_API_ID'])
    except ValueError:
        print(f"ERROR: TELEGRAM_API_ID must be a number, got: {required['TELEGRAM_API_ID']}")
        return False

    # Check optional Google Sheets config
    if not os.environ.get('GOOGLE_SHEETS_CREDENTIALS'):
        print("Warning: GOOGLE_SHEETS_CREDENTIALS not set - resume screening will use basic mode")
    if not os.environ.get('GOOGLE_SHEETS_ID'):
        print("Warning: GOOGLE_SHEETS_ID not set - resume screening will use basic mode")

    return True


def init_google_sheets():
    """Initialize Google Sheets client."""
    global gsheets_client

    creds_json = os.environ.get('GOOGLE_SHEETS_CREDENTIALS')
    if not creds_json:
        print("Google Sheets credentials not configured")
        return None

    try:
        creds_dict = json.loads(creds_json)
        scopes = [
            'https://www.googleapis.com/auth/spreadsheets.readonly',
            'https://www.googleapis.com/auth/drive.readonly'
        ]
        credentials = Credentials.from_service_account_info(creds_dict, scopes=scopes)
        gsheets_client = gspread.authorize(credentials)
        print("Google Sheets client initialized")
        return gsheets_client
    except Exception as e:
        print(f"Error initializing Google Sheets: {e}")
        return None


def get_message_delay_settings() -> tuple:
    """Get message delay settings from knowledgebase."""
    from shared.knowledgebase import COMMUNICATION_STYLE

    # Get delay setting from CRM configuration
    crm_settings = COMMUNICATION_STYLE.get('crm_settings', {})
    delay_setting = crm_settings.get('message_delay', 'normal')

    # Map delay setting to actual delay ranges (min, max)
    delay_map = {
        'instant': (0.0, 0.0),
        'fast': (0.5, 1.0),
        'normal': (1.5, 3.0),
        'slow': (3.0, 5.0),
        'very_slow': (5.0, 8.0),
    }

    return delay_map.get(delay_setting, (1.5, 3.0))


async def send_telegram_messages(event, telegram_client, message: str):
    """Send message(s) - splits on '---' and adds realistic typing delays."""
    parts = [part.strip() for part in message.split('---') if part.strip()]

    # Get delay settings from knowledgebase
    delay_min, delay_max = get_message_delay_settings()

    for i, part in enumerate(parts):
        # Show typing action while "thinking" and "typing"
        if delay_max > 0:
            async with telegram_client.action(event.chat_id, 'typing'):
                # Natural "thinking" delay based on settings
                thinking_delay = random.uniform(delay_min, delay_max)
                # Typing delay: ~0.03s per character (simulates typing speed)
                typing_delay = len(part) * 0.03
                # Total delay, capped at 10 seconds
                total_delay = min(thinking_delay + typing_delay, 10.0)
                await asyncio.sleep(total_delay)

        # First message replies directly to user's message, rest are normal responses
        if i == 0:
            await event.reply(part)
        else:
            await event.respond(part)

        # Add delay before next message (except for last one)
        if i < len(parts) - 1 and delay_max > 0:
            # Pause between messages based on settings
            between_delay = random.uniform(delay_min, delay_max)
            await asyncio.sleep(between_delay)


def setup_handlers(telegram_client):
    """Setup message handlers for the Telegram client."""

    @telegram_client.on(events.NewMessage(incoming=True))
    async def handle_message(event):
        # Skip if this is a file message (handled separately)
        if event.file:
            return

        # Only respond to private messages (ignore groups/channels)
        if not event.is_private:
            return

        # Check operating hours (8:30 AM - 10:00 PM Singapore time)
        if not is_within_operating_hours():
            print(f"Outside operating hours - not responding")
            return

        sender = await event.get_sender()
        user_id = sender.id
        username = sender.username or ""
        full_name = f"{sender.first_name or ''} {sender.last_name or ''}".strip()

        # Run spam protection checks
        if await check_spam_protection(event, user_id, username, event.text or ""):
            return

        print(f"Message from {full_name} (@{username}): {event.text}")

        # Check for training commands first (admin only)
        text = event.text or ""
        training_response = await handle_training_message(user_id, username, text)
        if training_response:
            # Training mode - send response directly (no delays for admin commands)
            await event.respond(training_response)
            return

        # Normal conversation mode
        async with telegram_client.action(event.chat_id, 'typing'):
            # Pass sender's name for context (especially for new conversations)
            response = await get_ai_response(user_id, text, candidate_name=full_name or username or None)

        # Send response with message splitting and delays
        await send_telegram_messages(event, telegram_client, response)
        # Note: Only create candidate record when resume is received (not on text messages)

    @telegram_client.on(events.NewMessage(incoming=True, func=lambda e: e.file))
    async def handle_file(event):
        # Only respond to private messages (ignore groups/channels)
        if not event.is_private:
            return

        # Check operating hours (8:30 AM - 10:00 PM Singapore time)
        if not is_within_operating_hours():
            print(f"Outside operating hours - not responding to file")
            return

        sender = await event.get_sender()
        user_id = sender.id
        username = sender.username or ""
        full_name = f"{sender.first_name or ''} {sender.last_name or ''}".strip()

        # Run spam protection checks (no text for file messages)
        if await check_spam_protection(event, user_id, username):
            return

        # Get file info
        file_name = event.file.name or "unknown"
        file_size = event.file.size or 0
        mime_type = event.file.mime_type or ""

        print(f"File received from {full_name} (@{username}): {file_name} ({mime_type}, {file_size} bytes)")

        # Check if it's a PDF or document
        is_resume = (
            mime_type == "application/pdf" or
            file_name.lower().endswith('.pdf') or
            mime_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"] or
            file_name.lower().endswith(('.doc', '.docx'))
        )

        if is_resume:
            # Restore conversation from database if this is a returning user
            await restore_conversation_from_db(user_id)

            # Store sender name if not already set
            current_state = get_conversation_state(user_id)
            if (full_name or username) and not current_state.get("candidate_name"):
                await update_conversation_state_async(user_id, candidate_name=full_name or username)

            await event.respond("thanks! will check it out")

            async with telegram_client.action(event.chat_id, 'typing'):
                # Download the file
                try:
                    file_bytes = await event.download_media(file=bytes)

                    if file_bytes:
                        # Extract text from resume and prepare for upload
                        upload_bytes = file_bytes
                        upload_name = file_name

                        if mime_type == "application/pdf" or file_name.lower().endswith('.pdf'):
                            resume_text = extract_text_from_pdf(file_bytes)
                            # If PyPDF2 extraction failed (image-based PDF like Canva), use Claude vision
                            if not resume_text or len(resume_text) < 100:
                                print(f"PyPDF2 extracted only {len(resume_text)} chars, trying vision API fallback...")
                                resume_text = await extract_text_from_pdf_with_vision(file_bytes)
                        elif mime_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"] or file_name.lower().endswith(('.doc', '.docx')):
                            resume_text = extract_text_from_word(file_bytes)
                            # Convert Word to PDF for preview compatibility
                            pdf_bytes = convert_word_to_pdf(file_bytes)
                            if pdf_bytes:
                                upload_bytes = pdf_bytes
                                # Change extension to .pdf
                                upload_name = os.path.splitext(file_name)[0] + '.pdf'
                                print(f"Converted Word doc to PDF: {upload_name}")
                        else:
                            # For other document types, just note that it was received
                            resume_text = f"[Document received: {file_name}]"

                        if resume_text and len(resume_text) > 100:
                            print(f"Extracted {len(resume_text)} characters from resume")

                            # Screen the resume FIRST to get candidate name
                            screening_result = await screen_resume(resume_text)
                            print(f"Screening result: {screening_result.get('recommendation', 'Unknown')}")

                            # Get candidate name for file naming
                            candidate_name = screening_result.get('candidate_name', full_name or 'Unknown')

                            # Create filename with candidate name
                            safe_name = "".join(c for c in candidate_name if c.isalnum() or c in (' ', '-', '_')).strip()
                            safe_name = safe_name.replace(' ', '_') if safe_name else 'Unknown'
                            final_upload_name = f"{safe_name}_Resume.pdf"

                            # Upload resume to storage with candidate name
                            resume_url = await upload_resume_to_storage(upload_bytes, final_upload_name, user_id)

                            # Save candidate with screening results and resume URL
                            await save_candidate(user_id, username, full_name, screening_result, resume_url)

                            # Update conversation state
                            matched_job = screening_result.get('job_matched', 'our open positions')
                            first_name = candidate_name.split()[0] if candidate_name else 'there'

                            await update_conversation_state_async(
                                user_id,
                                resume_received=True,
                                # Don't overwrite candidate_name - keep original Telegram display name
                                applied_role=matched_job,
                                stage="resume_received"
                            )

                            # Generate natural response using knowledgebase
                            # Identify role key for appropriate experience question
                            role_key = identify_role_from_text(matched_job)
                            response = get_resume_acknowledgment(first_name, role_key)

                            # Mark experience as discussed since we're asking about it
                            await update_conversation_state_async(user_id, experience_discussed=True, stage="experience_asked")

                            await event.respond(response)
                        else:
                            print("Could not extract sufficient text from resume")
                            await event.respond(
                                "thanks for ur resume! had a bit of trouble reading it but our team will review it manually. anything else i can help u with?"
                            )
                            # Note: Don't create candidate without successful resume processing
                    else:
                        print("Failed to download file")
                        await event.respond(
                            "had trouble downloading ur file. could u try sending it again?"
                        )
                except Exception as e:
                    print(f"Error processing file: {e}")
                    await event.respond(
                        "had some trouble processing ur file. our team will follow up with u. anything else i can help with?"
                    )
                    # Note: Don't create candidate on processing errors
        else:
            # Non-resume file - just respond, don't create candidate
            async with telegram_client.action(event.chat_id, 'typing'):
                response = await get_ai_response(user_id, f"[User sent a file: {file_name}]")
            await event.respond(response)


async def main():
    global client, anthropic_client, supabase_client

    print("=" * 50)
    print("Telegram Bot Starting...")
    print("=" * 50)

    # Validate environment variables
    print("Checking environment variables...")
    if not validate_env_vars():
        sys.exit(1)
    print("Environment variables OK")

    # Get environment variables
    api_id = int(os.environ.get('TELEGRAM_API_ID'))
    api_hash = os.environ.get('TELEGRAM_API_HASH')
    session_string = os.environ.get('TELEGRAM_SESSION_STRING')
    claude_api_key = os.environ.get('CLAUDE_API_KEY')
    supabase_url = os.environ.get('SUPABASE_URL')
    supabase_key = os.environ.get('SUPABASE_ANON_KEY')

    # Initialize Anthropic client
    print("Initializing Anthropic client...")
    try:
        anthropic_client = Anthropic(api_key=claude_api_key)
        print("Anthropic client OK")
    except Exception as e:
        print(f"ERROR: Failed to initialize Anthropic client: {e}")
        sys.exit(1)

    # Initialize Supabase client
    print("Initializing Supabase client...")
    try:
        supabase_client = create_client(supabase_url, supabase_key)
        print("Supabase client OK")
    except Exception as e:
        print(f"ERROR: Failed to initialize Supabase client: {e}")
        sys.exit(1)

    # Initialize Google Sheets client (optional)
    print("Initializing Google Sheets client...")
    init_google_sheets()

    # Initialize training system (admin users)
    print("Initializing training system...")
    init_admin_users()

    # Load knowledgebase from database
    print("Loading knowledgebase from database...")
    try:
        kb_loaded = await reload_from_database()
        if kb_loaded:
            print("Knowledgebase loaded from database")
        else:
            print("Using static knowledgebase (no DB entries found)")
    except Exception as e:
        print(f"Warning: Could not load knowledgebase from DB: {e}")
        print("Using static knowledgebase")

    # Generate embeddings for knowledgebase entries (RAG)
    if os.environ.get('OPENAI_API_KEY'):
        print("Checking knowledgebase embeddings (RAG)...")
        try:
            stats = await embed_existing_knowledge()
            if stats.get('processed', 0) > 0:
                print(f"Generated {stats['processed']} new embeddings")
            else:
                print("All entries already have embeddings")
        except Exception as e:
            print(f"Warning: Embedding generation skipped: {e}")
    else:
        print("Skipping RAG embeddings (OPENAI_API_KEY not set)")

    # Initialize Telegram client
    print("Initializing Telegram client...")
    print(f"  API ID: {api_id}")
    print(f"  Session string length: {len(session_string)} chars")

    try:
        client = TelegramClient(
            StringSession(session_string),
            api_id,
            api_hash,
            connection_retries=3,
            retry_delay=1
        )
        print("Telegram client created")
    except Exception as e:
        print(f"ERROR: Failed to create Telegram client: {e}")
        sys.exit(1)

    # Setup message handlers
    print("Setting up message handlers...")
    setup_handlers(client)
    print("Handlers OK")

    # Connect to Telegram
    print("Connecting to Telegram...")
    try:
        await asyncio.wait_for(client.connect(), timeout=30)
        print("Connected to Telegram")
    except asyncio.TimeoutError:
        print("ERROR: Connection to Telegram timed out after 30 seconds")
        sys.exit(1)
    except Exception as e:
        print(f"ERROR: Failed to connect to Telegram: {e}")
        sys.exit(1)

    # Check authorization
    print("Checking authorization...")
    try:
        if not await client.is_user_authorized():
            print("ERROR: Session is not authorized. Please regenerate your session string.")
            sys.exit(1)
        print("Authorization OK")
    except Exception as e:
        print(f"ERROR: Authorization check failed: {e}")
        sys.exit(1)

    # Get user info
    try:
        me = await client.get_me()
        print(f"Logged in as: {me.first_name} (@{me.username})")
    except Exception as e:
        print(f"Warning: Could not get user info: {e}")

    print("=" * 50)
    print("Bot is running! Waiting for messages...")
    print(f"Knowledgebase will auto-refresh every {KB_REFRESH_INTERVAL} seconds")
    print("=" * 50)

    # Start background task for periodic knowledgebase refresh
    asyncio.create_task(periodic_knowledgebase_refresh())

    await client.run_until_disconnected()


if __name__ == "__main__":
    asyncio.run(main())
