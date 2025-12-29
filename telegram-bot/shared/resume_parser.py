"""Resume text extraction utilities."""
from io import BytesIO
import os
import subprocess
import tempfile
import base64
import PyPDF2
from docx import Document


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text content from a PDF file using PyPDF2.

    Note: This only works for PDFs with selectable text. For image-based PDFs
    (like Canva resumes), use extract_text_from_pdf_with_vision() as a fallback.
    """
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return ""


async def extract_text_from_pdf_with_vision(pdf_bytes: bytes, anthropic_client=None) -> str:
    """Extract text from a PDF using Claude's vision API (for image-based PDFs like Canva resumes).

    This is used as a fallback when PyPDF2 cannot extract text (e.g., when text is rendered as images).

    Args:
        pdf_bytes: The PDF file as bytes
        anthropic_client: Optional Anthropic client. If not provided, will try to import from ai_screening.

    Returns:
        Extracted text from the PDF, or empty string on failure.
    """
    # Get anthropic client if not provided
    if anthropic_client is None:
        try:
            from shared.ai_screening import anthropic_client as shared_client
            anthropic_client = shared_client
        except ImportError:
            try:
                from ai_screening import anthropic_client as shared_client
                anthropic_client = shared_client
            except ImportError:
                pass

    if not anthropic_client:
        print("Error: Anthropic client not available for vision extraction")
        return ""

    try:
        # Encode PDF as base64
        pdf_base64 = base64.standard_b64encode(pdf_bytes).decode("utf-8")

        print("Using Claude vision API to extract text from image-based PDF...")

        response = anthropic_client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "document",
                            "source": {
                                "type": "base64",
                                "media_type": "application/pdf",
                                "data": pdf_base64
                            }
                        },
                        {
                            "type": "text",
                            "text": """Extract ALL text content from this resume/CV document.
Include everything: name, contact details, work experience, education, skills, certifications, etc.
Format it in a readable way, preserving the structure and sections.
Just output the extracted text, no commentary."""
                        }
                    ]
                }
            ]
        )

        extracted_text = response.content[0].text
        print(f"Vision API extracted {len(extracted_text)} characters from PDF")
        return extracted_text.strip()

    except Exception as e:
        print(f"Error extracting PDF text with vision API: {e}")
        return ""


def extract_text_from_word(doc_bytes: bytes) -> str:
    """Extract text content from a Word document (.docx)."""
    try:
        doc = Document(BytesIO(doc_bytes))
        text_parts = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"Error extracting Word document text: {e}")
        return ""


def convert_word_to_pdf(doc_bytes: bytes) -> bytes:
    """Convert a Word document to PDF using LibreOffice."""
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Write Word doc to temp file
            input_path = os.path.join(tmp_dir, "input.docx")
            with open(input_path, "wb") as f:
                f.write(doc_bytes)

            # Convert using LibreOffice
            subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", tmp_dir,
                input_path
            ], check=True, capture_output=True, timeout=60)

            # Read the converted PDF
            output_path = os.path.join(tmp_dir, "input.pdf")
            with open(output_path, "rb") as f:
                return f.read()
    except subprocess.TimeoutExpired:
        print("Error: Word to PDF conversion timed out")
        return None
    except Exception as e:
        print(f"Error converting Word to PDF: {e}")
        return None
